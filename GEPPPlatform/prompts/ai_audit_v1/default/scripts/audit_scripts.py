"""
Default AI Audit Script
Main entry point for the default audit system using OpenRouter LLM.
Processes transaction_audit_history records with 2-level threading.
"""

import logging
import time
import json
import traceback
from concurrent.futures import ThreadPoolExecutor, as_completed
from datetime import datetime, timezone, timedelta
from typing import Dict, Any, List, Optional, Callable
from pathlib import Path

from sqlalchemy.orm import Session
from sqlalchemy.orm.attributes import flag_modified
from sqlalchemy import and_

logger = logging.getLogger(__name__)

# Version marker for deployment verification
AUDIT_SCRIPT_VERSION = '2026-03-19-v5-per-record-evidence-only'

# Constants
ALLOWED_TRANSACTION_METHODS = ('origin', 'qr_input')
SETTINGS_PATH = Path(__file__).parent.parent / 'settings.json'

SUPPORTED_IMAGE_TYPES = {'image/jpeg', 'image/jpg', 'image/png', 'image/webp'}
SPREADSHEET_TYPES = {
    'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',  # .xlsx
    'application/vnd.ms-excel',  # .xls
    'text/csv',
    'application/csv',
}
PDF_TYPES = {'application/pdf'}
WORD_TYPES = {
    'application/vnd.openxmlformats-officedocument.wordprocessingml.document',  # .docx
    'application/msword',  # .doc
}
# All file types that are processed (for logging/tracking)
ALL_SUPPORTED_TYPES = SUPPORTED_IMAGE_TYPES | SPREADSHEET_TYPES | PDF_TYPES | WORD_TYPES


def _load_settings() -> Dict[str, Any]:
    with open(SETTINGS_PATH, 'r') as f:
        return json.load(f)


def _get_or_create_quota(db: Any, organization_id: int, subscription: Any) -> Any:
    """Get or create a SubscriptionMonthlyQuota for the current period."""
    from GEPPPlatform.models.subscriptions.subscription_monthly_quotas import SubscriptionMonthlyQuota

    duration_type = subscription.duration_type or 'monthly'
    now = datetime.now(timezone.utc)
    scope = now.strftime('%Y-%m') if duration_type == 'monthly' else now.strftime('%Y')

    quota = db.query(SubscriptionMonthlyQuota).filter(
        SubscriptionMonthlyQuota.organization_id == organization_id,
        SubscriptionMonthlyQuota.duration_type == duration_type,
        SubscriptionMonthlyQuota.duration_scope == scope,
        SubscriptionMonthlyQuota.deleted_date.is_(None),
    ).first()

    if not quota:
        quota = SubscriptionMonthlyQuota(
            organization_id=organization_id,
            duration_type=duration_type,
            duration_scope=scope,
            ai_audit_limit=subscription.ai_audit_limit or 10,
            ai_audit_usage=0,
            create_transaction_limit=subscription.create_transaction_limit or 100,
            create_transaction_usage=0,
        )
        db.add(quota)
        db.commit()

    return quota


def run_default_audit(get_session_factory: Callable) -> Dict[str, Any]:
    """
    Main entry point called from audit_cron.py.
    Manages Level-1 threading (one thread per org's audit_history record).

    Args:
        get_session_factory: Callable that returns a session factory (scoped_session)

    Returns:
        Summary dict of processing results
    """
    from GEPPPlatform.models.logs.transaction_audit_history import TransactionAuditHistory

    logger.info("Starting default audit processing")
    start_time = time.time()
    session_factory = get_session_factory()

    try:
        # Get a session for the initial query
        db = session_factory()
        try:
            # Step 1: Query pending/in_progress audit history, sorted by id ASC
            histories = db.query(TransactionAuditHistory).filter(
                and_(
                    TransactionAuditHistory.status.in_(['pending', 'in_progress']),
                    TransactionAuditHistory.deleted_date.is_(None)
                )
            ).order_by(TransactionAuditHistory.id.asc()).all()

            if not histories:
                logger.info("No pending/in_progress audit history records found")
                return {'success': True, 'message': 'No audit tasks to process', 'processed': 0}

            # Deduplicate: keep only the first (lowest id) per unique organization
            seen_orgs = set()
            unique_histories = []
            for h in histories:
                if h.organization_id not in seen_orgs:
                    seen_orgs.add(h.organization_id)
                    unique_histories.append((h.id, h.organization_id))

            logger.info(f"Found {len(unique_histories)} unique org audit tasks from {len(histories)} total records")
        finally:
            db.close()

        # Level-1 threading: one thread per org's audit history
        settings = _load_settings()
        max_workers = settings.get('max_thread_workers', 10)
        results = []
        with ThreadPoolExecutor(max_workers=min(len(unique_histories), max_workers)) as executor:
            futures = {}
            for history_id, org_id in unique_histories:
                future = executor.submit(
                    _process_single_audit_history,
                    history_id, org_id, session_factory
                )
                futures[future] = (history_id, org_id)

            for future in as_completed(futures):
                history_id, org_id = futures[future]
                try:
                    result = future.result()
                    results.append(result)
                    logger.info(f"Audit history #{history_id} (org {org_id}) completed: {result.get('status')}")
                except Exception as e:
                    logger.error(f"Audit history #{history_id} (org {org_id}) failed: {str(e)}")
                    logger.error(traceback.format_exc())
                    results.append({'history_id': history_id, 'org_id': org_id, 'status': 'error', 'error': str(e)})

        elapsed = time.time() - start_time
        total_processed = sum(r.get('processed_count', 0) for r in results)
        logger.info(f"Default audit completed. {total_processed} transactions processed in {elapsed:.2f}s")

        return {
            'success': True,
            'message': f'Processed {total_processed} transactions across {len(results)} batches',
            'processed': total_processed,
            'batches': len(results),
            'elapsed_seconds': round(elapsed, 2),
            'results': results,
        }

    except Exception as e:
        logger.error(f"Error in run_default_audit: {str(e)}")
        logger.error(traceback.format_exc())
        return {'success': False, 'error': str(e)}


def _process_single_audit_history(
    history_id: int,
    organization_id: int,
    session_factory
) -> Dict[str, Any]:
    """
    Process one transaction_audit_history record.
    Steps 2-9 for a single org batch.
    """
    from GEPPPlatform.models.logs.transaction_audit_history import TransactionAuditHistory
    from GEPPPlatform.models.transactions.transactions import Transaction, AIAuditStatus
    from GEPPPlatform.models.transactions.transaction_records import TransactionRecord
    from GEPPPlatform.models.transactions.ai_audit_document_types import AiAuditDocumentType
    from GEPPPlatform.models.transactions.ai_audit_column_details import AiAuditColumnDetail
    from GEPPPlatform.models.subscriptions.organization_audit_settings import (
        OrganizationAuditDocRequireTypes, OrganizationAuditCheckColumns
    )
    from GEPPPlatform.models.subscriptions.subscription_models import Subscription
    from GEPPPlatform.models.subscriptions.subscription_monthly_quotas import SubscriptionMonthlyQuota
    from ..clients.llm_client import get_default_audit_llm

    db = session_factory()
    try:
        # Step 2: Get history and mark as in_progress
        history = db.query(TransactionAuditHistory).filter(
            TransactionAuditHistory.id == history_id
        ).first()

        if not history:
            return {'history_id': history_id, 'status': 'not_found', 'processed_count': 0}

        if history.status == 'pending':
            history.status = 'in_progress'
            history.started_at = datetime.now(timezone.utc)
            db.commit()

        # Step 3: Get queued transactions from the history's transaction list
        tx_ids = history.transactions or []
        if not tx_ids:
            history.status = 'completed'
            history.completed_at = datetime.now(timezone.utc)
            db.commit()
            return {'history_id': history_id, 'status': 'completed', 'processed_count': 0, 'message': 'No transactions in batch'}

        settings = _load_settings()
        max_tx_per_batch = settings.get('max_transactions_per_batch', 25)
        max_workers = settings.get('max_thread_workers', 10)

        transactions = db.query(Transaction).filter(
            and_(
                Transaction.id.in_(tx_ids),
                Transaction.ai_audit_status == AIAuditStatus.queued,
                Transaction.transaction_method.in_(ALLOWED_TRANSACTION_METHODS),
                Transaction.deleted_date.is_(None)
            )
        ).order_by(Transaction.id.asc()).limit(max_tx_per_batch).all()

        if not transactions:
            # Check if there are any remaining queued transactions
            remaining = db.query(Transaction).filter(
                and_(
                    Transaction.id.in_(tx_ids),
                    Transaction.ai_audit_status == AIAuditStatus.queued,
                    Transaction.deleted_date.is_(None)
                )
            ).count()

            if remaining == 0:
                history.status = 'completed'
                history.completed_at = datetime.now(timezone.utc)
                db.commit()

            return {'history_id': history_id, 'status': 'completed', 'processed_count': 0, 'message': 'No eligible queued transactions'}

        logger.info(f"Processing {len(transactions)} transactions for audit history #{history_id} (org {organization_id})")

        # Load org config
        doc_type_specs = db.query(AiAuditDocumentType).filter(
            AiAuditDocumentType.deleted_date.is_(None),
            AiAuditDocumentType.is_active.is_(True)
        ).all()
        doc_type_specs_list = [dt.to_dict() for dt in doc_type_specs]

        doc_requires = db.query(OrganizationAuditDocRequireTypes).filter(
            OrganizationAuditDocRequireTypes.organization_id == organization_id,
            OrganizationAuditDocRequireTypes.deleted_date.is_(None)
        ).first()

        check_columns = db.query(OrganizationAuditCheckColumns).filter(
            OrganizationAuditCheckColumns.organization_id == organization_id,
            OrganizationAuditCheckColumns.deleted_date.is_(None)
        ).first()

        # Load column details for resolving IDs → column names
        column_details = db.query(AiAuditColumnDetail).filter(
            AiAuditColumnDetail.deleted_date.is_(None),
            AiAuditColumnDetail.is_active.is_(True),
        ).all()
        column_detail_map = {cd.id: cd.column_name for cd in column_details}

        config = {
            'doc_type_specs': doc_type_specs_list,
            'doc_requires': doc_requires.to_dict() if doc_requires else {'transaction_document_requires': [], 'record_document_requires': []},
            'check_columns': check_columns.to_dict() if check_columns else {'transaction_checks': {}, 'transaction_record_checks': []},
            'column_detail_map': column_detail_map,
        }

        # Initialize LLM
        settings = _load_settings()
        # model_version = settings.get('model', 'x-ai/grok-4.1-fast')
        model_version = settings.get('model', 'google/gemini-3-flash-preview')
        # Quota check: count total records and verify against subscription quota
        total_record_count = 0
        for txn in transactions:
            rec_ids = txn.transaction_records or []
            total_record_count += len(rec_ids)

        subscription = db.query(Subscription).filter(
            Subscription.organization_id == organization_id,
            Subscription.status == 'active',
            Subscription.deleted_date.is_(None),
        ).order_by(Subscription.id.desc()).first()

        quota_id = None
        has_subscription = False
        if subscription:
            has_subscription = True
            quota = _get_or_create_quota(db, organization_id, subscription)
            quota_id = quota.id

            remaining = quota.ai_audit_limit - quota.ai_audit_usage
            if remaining < total_record_count and not (subscription.allow_ai_audit_exceed_quota):
                # Quota exhausted — skip this batch
                logger.warning(
                    f"Org {organization_id}: AI audit quota exhausted "
                    f"(remaining={remaining}, needed={total_record_count}, limit={quota.ai_audit_limit})"
                )
                history.error_message = (
                    f"โควต้า AI Audit หมด (เหลือ {remaining} จากทั้งหมด {quota.ai_audit_limit}, "
                    f"ต้องการ {total_record_count} records)"
                )
                db.commit()
                db.close()
                return {
                    'history_id': history_id,
                    'org_id': organization_id,
                    'status': 'quota_exhausted',
                    'processed_count': 0,
                    'message': f'AI audit quota exhausted (remaining={remaining}, needed={total_record_count})',
                }

        # Extract transaction IDs before closing session (ORM objects become detached after close)
        transaction_ids = [tx.id for tx in transactions]
        db.close()  # Close parent session before threading

        # Level-2 threading: one thread per transaction
        tx_results = []
        with ThreadPoolExecutor(max_workers=min(len(transaction_ids), max_workers)) as executor:
            futures = {}
            for tx_id in transaction_ids:
                future = executor.submit(
                    _process_single_transaction,
                    tx_id, organization_id, config, model_version, session_factory
                )
                futures[future] = tx_id

            for future in as_completed(futures):
                tx_id = futures[future]
                try:
                    result = future.result()
                    tx_results.append(result)
                except Exception as e:
                    logger.error(f"Transaction #{tx_id} audit failed: {str(e)}")
                    logger.error(traceback.format_exc())
                    tx_results.append({
                        'transaction_id': tx_id,
                        'status': 'failed',
                        'error': str(e),
                    })
                    # Mark the transaction as failed
                    _mark_transaction_failed(tx_id, str(e), session_factory)

        # Update audit history with results
        db = session_factory()
        try:
            history = db.query(TransactionAuditHistory).filter(
                TransactionAuditHistory.id == history_id
            ).first()

            if history:
                approved = sum(1 for r in tx_results if r.get('status') == 'approved')
                rejected = sum(1 for r in tx_results if r.get('status') == 'rejected')
                failed = sum(1 for r in tx_results if r.get('status') == 'failed')

                history.processed_transactions = (history.processed_transactions or 0) + len(tx_results)
                history.approved_count = (history.approved_count or 0) + approved
                history.rejected_count = (history.rejected_count or 0) + rejected

                # Collect audit_record IDs and append to history.audit_records
                new_audit_ids = [r['audit_record_id'] for r in tx_results if r.get('audit_record_id')]
                if new_audit_ids:
                    existing_ids = list(history.audit_records or [])
                    history.audit_records = existing_ids + new_audit_ids
                    flag_modified(history, 'audit_records')

                # Check if all transactions are processed
                remaining_queued = db.query(Transaction).filter(
                    and_(
                        Transaction.id.in_(tx_ids),
                        Transaction.ai_audit_status == AIAuditStatus.queued,
                        Transaction.deleted_date.is_(None)
                    )
                ).count()

                if remaining_queued == 0:
                    history.status = 'completed'
                    history.completed_at = datetime.now(timezone.utc)

                history.audit_info = {
                    'last_batch': {
                        'processed': len(tx_results),
                        'approved': approved,
                        'rejected': rejected,
                        'failed': failed,
                        'timestamp': datetime.now(timezone.utc).isoformat(),
                    }
                }
                flag_modified(history, 'audit_info')
                db.commit()

            # Update quota usage: count actually processed records (not failed)
            if has_subscription and quota_id:
                processed_tx_ids = [r['transaction_id'] for r in tx_results if r.get('status') in ('approved', 'rejected')]
                if processed_tx_ids:
                    processed_record_count = db.query(TransactionRecord).filter(
                        TransactionRecord.created_transaction_id.in_(processed_tx_ids),
                        TransactionRecord.deleted_date.is_(None),
                    ).count()

                    if processed_record_count > 0:
                        fresh_quota = db.query(SubscriptionMonthlyQuota).filter(
                            SubscriptionMonthlyQuota.id == quota_id
                        ).first()
                        if fresh_quota:
                            fresh_quota.ai_audit_usage = (fresh_quota.ai_audit_usage or 0) + processed_record_count
                            db.commit()
                            logger.info(f"Org {organization_id}: quota updated, ai_audit_usage += {processed_record_count}")

        finally:
            db.close()

        return {
            'history_id': history_id,
            'org_id': organization_id,
            'status': 'completed',
            'processed_count': len(tx_results),
            'approved': sum(1 for r in tx_results if r.get('status') == 'approved'),
            'rejected': sum(1 for r in tx_results if r.get('status') == 'rejected'),
            'failed': sum(1 for r in tx_results if r.get('status') == 'failed'),
        }

    except Exception as e:
        logger.error(f"Error processing audit history #{history_id}: {str(e)}")
        logger.error(traceback.format_exc())
        try:
            db = session_factory()
            history = db.query(TransactionAuditHistory).filter(
                TransactionAuditHistory.id == history_id
            ).first()
            if history:
                history.error_message = str(e)[:500]
                db.commit()
            db.close()
        except Exception:
            pass
        raise
    finally:
        try:
            db.close()
        except Exception:
            pass


def _mark_transaction_failed(tx_id: int, error_msg: str, session_factory) -> None:
    """Mark a transaction as failed audit"""
    from GEPPPlatform.models.transactions.transactions import Transaction, AIAuditStatus

    db = session_factory()
    try:
        tx = db.query(Transaction).filter(Transaction.id == tx_id).first()
        if tx:
            tx.ai_audit_status = AIAuditStatus.failed
            tx.ai_audit_note = {
                'status': 'failed',
                'error': error_msg[:500],
            }
            tx.ai_audit_date = datetime.now(timezone.utc)
            flag_modified(tx, 'ai_audit_note')
            db.commit()
    except Exception as e:
        logger.error(f"Failed to mark transaction #{tx_id} as failed: {str(e)}")
        db.rollback()
    finally:
        db.close()


def _process_single_transaction(
    transaction_id: int,
    organization_id: int,
    config: Dict[str, Any],
    model_version: str,
    session_factory
) -> Dict[str, Any]:
    """
    Process a single transaction through steps 4-9.
    Each call gets its own DB session.
    """
    from GEPPPlatform.models.transactions.transactions import Transaction, AIAuditStatus
    from GEPPPlatform.models.transactions.transaction_records import TransactionRecord
    from ..clients.llm_client import get_default_audit_llm

    processing_start = time.time()
    total_token_usage = {'input_tokens': 0, 'output_tokens': 0}
    print(f"[AUDIT-VERSION] Tx #{transaction_id}: script version={AUDIT_SCRIPT_VERSION}")

    db = session_factory()
    try:
        # Load transaction and its records
        tx = db.query(Transaction).filter(Transaction.id == transaction_id).first()
        if not tx:
            return {'transaction_id': transaction_id, 'status': 'failed', 'error': 'Transaction not found'}

        record_ids = tx.transaction_records or []
        records = []
        if record_ids:
            records = db.query(TransactionRecord).filter(
                TransactionRecord.id.in_(record_ids)
            ).all()

        llm = get_default_audit_llm()

        # === Step 4: Collect images ===
        t0 = time.time()
        image_data = _step4_collect_images(tx, records, db)
        logger.info(f"Tx #{transaction_id} Step 4 (collect files): {time.time()-t0:.1f}s — {len(image_data.get('all_files', {}))} files")

        # === Step 6: Classify evidence ===
        t0 = time.time()
        classified_evidence = _step6_classify_evidence(
            image_data, config['doc_type_specs'], llm, db, total_token_usage
        )
        logger.info(f"Tx #{transaction_id} Step 6 (classify): {time.time()-t0:.1f}s — {len(classified_evidence)} files classified")

        # === Step 7: Check required docs ===
        all_record_ids = [r.id for r in records]
        doc_check = _step7_check_required_docs(
            classified_evidence, config['doc_requires'],
            image_data.get('transaction_file_ids', []),
            image_data.get('record_file_ids', {}),
            all_record_ids=all_record_ids,
            doc_type_specs=config['doc_type_specs'],
        )

        # === Step 8: Checklist-based matching (Phase A → B → C) ===
        rec_checks = config['check_columns'].get('transaction_record_checks', [])
        column_detail_map = config.get('column_detail_map', {})
        # transaction_record_checks is a list of ai_audit_column_details IDs, e.g. [6, 7, 8]
        if isinstance(rec_checks, list):
            checklist_columns = [column_detail_map[cid] for cid in rec_checks if cid in column_detail_map]
        else:
            # Legacy fallback: old format {"column_name": true/false}
            checklist_columns = [k for k, v in rec_checks.items() if v]

        # Resolve names in batch
        names = _resolve_names_batch(tx, records, db)
        all_records_data = [_build_record_data(r, names) for r in records]

        # Phase A: Transaction-level evidence check
        t0 = time.time()
        tx_checklist = _step8a_transaction_level_check(
            image_data, classified_evidence, checklist_columns, all_records_data,
            llm, total_token_usage, transaction_id,
        )
        logger.info(f"Tx #{transaction_id} Step 8A (tx-level): {time.time()-t0:.1f}s")

        # Phase B: Record-level evidence check (only unmatched columns)
        unmatched_columns = [
            col for col in checklist_columns
            if not (tx_checklist[col]['match'] and tx_checklist[col]['found'])
        ]

        per_record_results = {}
        if unmatched_columns and records:
            t0 = time.time()
            per_record_results = _step8b_record_level_check(
                records, image_data, classified_evidence, unmatched_columns,
                names, llm, total_token_usage, transaction_id,
            )
            logger.info(f"Tx #{transaction_id} Step 8B (record-level): {time.time()-t0:.1f}s — {len(unmatched_columns)} unmatched cols")
        else:
            logger.info(f"Tx #{transaction_id} Step 8B skipped — all columns matched at tx-level")

        # Phase C: Determine final status
        required_columns = checklist_columns  # all checked columns are required
        final_determination = _determine_final_status(
            tx_checklist, per_record_results, checklist_columns, required_columns, doc_check,
        )

        # === Step 9: Compose & Save ===
        t0 = time.time()
        final_result = _step9_compose_and_save(
            tx, records, doc_check, final_determination, classified_evidence,
            per_record_results, tx_checklist, total_token_usage,
            processing_start, organization_id, model_version, db, llm,
            image_data=image_data,
        )
        logger.info(f"Tx #{transaction_id} Step 9 (compose): {time.time()-t0:.1f}s")

        db.commit()
        return final_result

    except Exception as e:
        db.rollback()
        logger.error(f"Error processing transaction #{transaction_id}: {str(e)}")
        logger.error(traceback.format_exc())
        raise
    finally:
        db.close()


def _step4_collect_images(
    tx: Any,
    records: List[Any],
    db: Session
) -> Dict[str, Any]:
    """
    Step 4: Collect image file IDs from transaction and records.
    Filter: skip strings (URLs not allowed in default), keep integers (file IDs).
    Query File table and generate presigned read URLs.
    """
    from GEPPPlatform.models.cores.files import File
    from GEPPPlatform.services.cores.transactions.presigned_url_service import TransactionPresignedUrlService

    result = {
        'transaction_images': [],
        'record_images': {},
        'transaction_file_ids': [],
        'record_file_ids': {},
        'all_files': {},  # file_id -> {presigned_url, mime_type, s3_key, source}
    }

    # Collect file IDs from transaction
    tx_image_ids = []
    for img in (tx.images or []):
        if isinstance(img, int):
            tx_image_ids.append(img)

    # Collect file IDs from each record
    record_image_ids = {}
    for record in records:
        rec_ids = []
        for img in (record.images or []):
            if isinstance(img, int):
                rec_ids.append(img)
        if rec_ids:
            record_image_ids[record.id] = rec_ids

    # Combine all file IDs for batch query
    all_file_ids = set(tx_image_ids)
    for ids in record_image_ids.values():
        all_file_ids.update(ids)

    if not all_file_ids:
        logger.info(f"Transaction #{tx.id}: No valid file IDs found in images")
        return result

    # Query files and generate presigned URLs
    files = db.query(File).filter(File.id.in_(all_file_ids)).all()
    file_map = {f.id: f for f in files}

    try:
        presigned_service = TransactionPresignedUrlService()
        presigned_result = presigned_service.get_transaction_file_view_presigned_urls_by_ids(
            file_ids=list(all_file_ids),
            db=db,
            organization_id=tx.organization_id,
            user_id=tx.created_by_id or 0,
            expiration_seconds=7200
        )

        presigned_urls = presigned_result.get('presigned_urls', {}) if presigned_result.get('success') else {}
    except Exception as e:
        logger.error(f"Failed to generate presigned URLs: {str(e)}")
        presigned_urls = {}

    # Build result
    for file_id in all_file_ids:
        f = file_map.get(file_id)
        url_data = presigned_urls.get(file_id, {})
        if f and url_data.get('view_url'):
            result['all_files'][file_id] = {
                'file_id': file_id,
                'presigned_url': url_data['view_url'],
                'mime_type': f.mime_type,
                's3_key': f.s3_key,
            }

    # Exclude files that belong to specific records from tx-level
    # If a file_id appears in both tx.images and record.images, it's record-level only
    all_record_fids = set()
    for fids in record_image_ids.values():
        all_record_fids.update(fids)
    pure_tx_file_ids = [fid for fid in tx_image_ids if fid not in all_record_fids]

    result['transaction_file_ids'] = pure_tx_file_ids
    result['transaction_images'] = [result['all_files'][fid] for fid in pure_tx_file_ids if fid in result['all_files']]
    result['record_file_ids'] = record_image_ids
    for rec_id, fids in record_image_ids.items():
        result['record_images'][rec_id] = [result['all_files'][fid] for fid in fids if fid in result['all_files']]

    logger.info(f"Transaction #{tx.id}: Collected {len(result['all_files'])} files ({len(result['transaction_images'])} tx-level, {sum(len(v) for v in result['record_images'].values())} record-level)")
    return result


def _download_file_bytes(presigned_url: str, timeout: int = 30) -> bytes:
    """Download file content from a presigned URL."""
    from urllib.request import urlopen
    response = urlopen(presigned_url, timeout=timeout)
    return response.read()


def _parse_spreadsheet_to_text(file_bytes: bytes, mime_type: str) -> str:
    """
    Parse spreadsheet (xlsx/xls/csv) to text.
    Uses openpyxl for xlsx, csv stdlib for csv.
    Returns formatted text representation of all sheets/data.
    """
    import io
    import csv

    if mime_type in ('text/csv', 'application/csv'):
        text_content = file_bytes.decode('utf-8', errors='replace')
        reader = csv.reader(io.StringIO(text_content))
        rows = list(reader)
        if not rows:
            return "(Empty CSV file)"
        lines = []
        for i, row in enumerate(rows):
            lines.append(f"Row {i+1}: {' | '.join(row)}")
        return f"[CSV — {len(rows)} rows]\n" + '\n'.join(lines)

    else:
        # Excel (xlsx/xls)
        try:
            import openpyxl
            wb = openpyxl.load_workbook(io.BytesIO(file_bytes), read_only=True, data_only=True)
        except Exception as e:
            logger.error(f"Failed to parse Excel file: {e}")
            return f"(Failed to parse Excel file: {str(e)})"

        all_sheets_text = []
        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            rows = []
            for row in ws.iter_rows(values_only=True):
                cell_values = [str(c) if c is not None else '' for c in row]
                if any(v.strip() for v in cell_values):
                    rows.append(cell_values)

            if not rows:
                all_sheets_text.append(f"[Sheet: {sheet_name}] (Empty)")
                continue

            lines = []
            for i, row in enumerate(rows):
                lines.append(f"Row {i+1}: {' | '.join(row)}")
            all_sheets_text.append(f"[Sheet: {sheet_name} — {len(rows)} rows]\n" + '\n'.join(lines))

        wb.close()
        return '\n\n'.join(all_sheets_text)


def _parse_pdf_to_text(file_bytes: bytes) -> str:
    """
    Parse PDF to text using pypdf (pure Python, no native deps).
    Returns all pages as text.
    """
    import io
    try:
        from pypdf import PdfReader
    except ImportError:
        # Fallback: try PyPDF2 (older name)
        try:
            from PyPDF2 import PdfReader
        except ImportError:
            logger.error("Neither pypdf nor PyPDF2 available for PDF parsing")
            return "(PDF parsing unavailable — pypdf not installed)"

    try:
        reader = PdfReader(io.BytesIO(file_bytes))
        pages_text = []
        for i, page in enumerate(reader.pages):
            text = page.extract_text() or ''
            if text.strip():
                pages_text.append(f"[Page {i+1}]\n{text.strip()}")
        if not pages_text:
            return "(PDF has no extractable text — may be scanned/image-based)"
        return '\n\n'.join(pages_text)
    except Exception as e:
        logger.error(f"Failed to parse PDF: {e}")
        return f"(Failed to parse PDF: {str(e)})"


def _parse_word_to_text(file_bytes: bytes) -> str:
    """
    Parse Word document (.docx) to text.
    Uses python-docx for .docx files.
    """
    import io
    try:
        from docx import Document
    except ImportError:
        logger.error("python-docx not available for Word parsing")
        return "(Word parsing unavailable — python-docx not installed)"

    try:
        doc = Document(io.BytesIO(file_bytes))
        paragraphs = []
        for para in doc.paragraphs:
            if para.text.strip():
                paragraphs.append(para.text.strip())

        # Also extract tables
        for i, table in enumerate(doc.tables):
            table_rows = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                table_rows.append(' | '.join(cells))
            if table_rows:
                paragraphs.append(f"[Table {i+1}]\n" + '\n'.join(table_rows))

        if not paragraphs:
            return "(Empty Word document)"
        return '\n\n'.join(paragraphs)
    except Exception as e:
        logger.error(f"Failed to parse Word document: {e}")
        return f"(Failed to parse Word document: {str(e)})"


def _step6_classify_evidence(
    image_data: Dict[str, Any],
    doc_type_specs: List[Dict],
    llm: Any,
    db: Session,
    token_usage: Dict[str, int]
) -> List[Dict[str, Any]]:
    """
    Step 6: Classify each evidence file and extract data using LLM.
    Store results in File.observation.
    """
    from GEPPPlatform.models.cores.files import File
    import yaml

    all_files = image_data.get('all_files', {})
    if not all_files:
        return []

    # Build document types spec string for the prompt
    types_for_prompt = []
    for dt in doc_type_specs:
        types_for_prompt.append({
            'id': dt['id'],
            'name_en': dt['name_en'],
            'name_th': dt['name_th'],
            'description': dt.get('description_en') or dt.get('description', ''),
            'extract_list': dt['extract_list'],
        })
    doc_types_spec_str = json.dumps(types_for_prompt, ensure_ascii=False, indent=2)

    # Load prompt template
    prompt_path = Path(__file__).parent.parent / 'prompts' / 'templates' / 'evidence_classify.yaml'
    with open(prompt_path, 'r') as f:
        prompt_data = yaml.safe_load(f)
    prompt_text = prompt_data['template'].format(document_types_spec=doc_types_spec_str)

    classified = []

    # Classify each file (threaded)
    def classify_single_file(file_id, file_info):
        from ..clients.llm_client import call_llm_with_images, call_llm_text_only, parse_json_response
        try:
            mime = (file_info.get('mime_type') or '').lower()
            presigned_url = file_info['presigned_url']

            combined_prompt = None  # Will be set for non-image types

            if mime in SUPPORTED_IMAGE_TYPES:
                # Image: send as image_url to LLM
                response = call_llm_with_images(llm, prompt_text, [presigned_url])

            elif mime in SPREADSHEET_TYPES:
                # Spreadsheet: download, parse all sheets to text, send as text-only
                file_bytes = _download_file_bytes(presigned_url)
                parsed_text = _parse_spreadsheet_to_text(file_bytes, mime)
                combined_prompt = (
                    f"{prompt_text}\n\n"
                    f"--- SPREADSHEET CONTENT (from uploaded file) ---\n"
                    f"{parsed_text}\n"
                    f"--- END SPREADSHEET CONTENT ---"
                )
                response = call_llm_text_only(llm, combined_prompt)

            elif mime in PDF_TYPES:
                # PDF: download, extract text from all pages, send as text-only
                file_bytes = _download_file_bytes(presigned_url)
                parsed_text = _parse_pdf_to_text(file_bytes)
                combined_prompt = (
                    f"{prompt_text}\n\n"
                    f"--- PDF CONTENT (from uploaded file) ---\n"
                    f"{parsed_text}\n"
                    f"--- END PDF CONTENT ---"
                )
                response = call_llm_text_only(llm, combined_prompt)

            elif mime in WORD_TYPES:
                # Word document: download, extract text + tables, send as text-only
                file_bytes = _download_file_bytes(presigned_url)
                parsed_text = _parse_word_to_text(file_bytes)
                combined_prompt = (
                    f"{prompt_text}\n\n"
                    f"--- WORD DOCUMENT CONTENT (from uploaded file) ---\n"
                    f"{parsed_text}\n"
                    f"--- END WORD DOCUMENT CONTENT ---"
                )
                response = call_llm_text_only(llm, combined_prompt)

            else:
                logger.warning(f"File #{file_id} has unsupported mime type: {mime}, skipping")
                return {
                    'file_id': file_id,
                    'document_type_id': 0,
                    'document_type_name': 'skipped_unsupported',
                    'extracted_data': {},
                    'confidence': 0.0,
                    'skipped_reason': f'Unsupported mime type: {mime}',
                }

            # Parse response with retry on JSON failure
            usage = response.get('usage', {})
            token_usage['input_tokens'] += usage.get('input_tokens', 0)
            token_usage['output_tokens'] += usage.get('output_tokens', 0)

            parse_error = None
            for attempt in range(3):
                try:
                    parsed = parse_json_response(response['content'])
                    return {
                        'file_id': file_id,
                        'document_type_id': parsed.get('document_type_id', 0),
                        'document_type_name': parsed.get('document_type_name', 'unknown'),
                        'extracted_data': parsed.get('extracted_data', {}),
                        'confidence': parsed.get('confidence', 0.0),
                    }
                except (json.JSONDecodeError, Exception) as je:
                    parse_error = je
                    if attempt < 2:
                        logger.warning(f"File #{file_id}: JSON parse failed (attempt {attempt+1}/3), retrying LLM call...")
                        try:
                            if combined_prompt:
                                response = call_llm_text_only(llm, combined_prompt)
                            else:
                                response = call_llm_with_images(llm, prompt_text, [presigned_url])
                            u = response.get('usage', {})
                            token_usage['input_tokens'] += u.get('input_tokens', 0)
                            token_usage['output_tokens'] += u.get('output_tokens', 0)
                        except Exception:
                            pass

            logger.error(f"File #{file_id}: JSON parse failed after 3 attempts: {str(parse_error)}")
            return {
                'file_id': file_id,
                'document_type_id': 0,
                'document_type_name': 'error',
                'extracted_data': {},
                'confidence': 0.0,
                'error': str(parse_error),
            }
        except Exception as e:
            logger.error(f"Failed to classify file #{file_id}: {str(e)}")
            return {
                'file_id': file_id,
                'document_type_id': 0,
                'document_type_name': 'error',
                'extracted_data': {},
                'confidence': 0.0,
                'error': str(e),
            }

    # Always re-classify all evidence files (no cache)
    file_ids_to_classify = list(all_files.keys())
    file_objects = {f.id: f for f in db.query(File).filter(File.id.in_(file_ids_to_classify)).all()}

    if False:
        logger.info(f"Step 6: {cached_count} files from cache, {len(file_ids_to_classify)} to classify")

    # Classify remaining files in parallel (LLM calls are thread-safe, DB writes are not)
    if file_ids_to_classify:
        new_results = []
        with ThreadPoolExecutor(max_workers=len(file_ids_to_classify)) as classify_executor:
            futures = {
                classify_executor.submit(classify_single_file, fid, all_files[fid]): fid
                for fid in file_ids_to_classify
            }
            for future in as_completed(futures):
                try:
                    result = future.result()
                    new_results.append(result)
                    classified.append(result)
                except Exception as e:
                    fid = futures[future]
                    logger.error(f"Classify file #{fid} thread failed: {str(e)}")

        # Write observations to DB sequentially
        for result in new_results:
            try:
                file_obj = file_objects.get(result['file_id'])
                if not file_obj:
                    file_obj = db.query(File).filter(File.id == result['file_id']).first()
                if file_obj:
                    observation = {
                        'document_type_id': result.get('document_type_id', 0),
                        'document_type_name': result.get('document_type_name', 'unknown'),
                        'extracted_data': result.get('extracted_data', {}),
                        'confidence': result.get('confidence', 0.0),
                        'classified_at': datetime.now(timezone.utc).isoformat(),
                    }
                    if result.get('skipped_reason'):
                        observation['skipped_reason'] = result['skipped_reason']
                    if result.get('error'):
                        observation['error'] = result['error']
                    file_obj.observation = observation
                    flag_modified(file_obj, 'observation')
            except Exception as e:
                logger.error(f"Failed to update File.observation for #{result['file_id']}: {str(e)}")

        # Commit all observation updates
        try:
            db.commit()
        except Exception as e:
            logger.error(f"Failed to commit File.observation updates: {str(e)}")

    logger.info(f"Classified {len(classified)} evidence files")
    return classified


def _step7_check_required_docs(
    classified_evidence: List[Dict],
    doc_requires: Dict[str, Any],
    transaction_file_ids: List[int],
    record_file_ids: Dict[int, List[int]],
    all_record_ids: List[int] = None,
    doc_type_specs: List[Dict] = None,
) -> Dict[str, Any]:
    """
    Step 7: Check if required document types are present at record level.
    Transaction images are reusable for records (counted toward record requirements).
    Checks ALL records (including those with no images) against requirements.
    Returns missing doc info with resolved names.
    """
    rec_doc_requires = doc_requires.get('record_document_requires', [])

    if not rec_doc_requires:
        return {'all_present': True, 'missing_transaction_docs': [], 'missing_record_docs': {}}

    # Build doc type ID → name map
    dt_name_map = {}
    if doc_type_specs:
        for dt in doc_type_specs:
            dt_id = dt.get('id')
            dt_name_map[dt_id] = {
                'id': dt_id,
                'name_en': dt.get('name_en', f'Document Type #{dt_id}'),
                'name_th': dt.get('name_th', f'ประเภทเอกสาร #{dt_id}'),
            }

    def _resolve_doc_name(dt_id):
        """Return {id, name_en, name_th} for a doc type ID."""
        if dt_id in dt_name_map:
            return dt_name_map[dt_id]
        return {'id': dt_id, 'name_en': f'Document Type #{dt_id}', 'name_th': f'ประเภทเอกสาร #{dt_id}'}

    # Map file_id to its classified document_type_id
    file_to_type = {}
    for ev in classified_evidence:
        if ev.get('document_type_id') and ev['document_type_id'] != 0:
            file_to_type[ev['file_id']] = ev['document_type_id']

    # Collect transaction-level present types (reusable for all records)
    tx_present_types = set()
    for fid in transaction_file_ids:
        if fid in file_to_type:
            tx_present_types.add(file_to_type[fid])

    # Check ALL records against requirements (not just those with files)
    # Use all_record_ids to ensure records with no images are also checked
    check_record_ids = set(all_record_ids or []) | set(record_file_ids.keys())

    missing_records = {}
    for rec_id in check_record_ids:
        rec_fids = record_file_ids.get(rec_id, [])
        rec_present_types = set(tx_present_types)  # Include tx-level images
        for fid in rec_fids:
            if fid in file_to_type:
                rec_present_types.add(file_to_type[fid])

        missing_rec = [_resolve_doc_name(dt_id) for dt_id in rec_doc_requires if dt_id not in rec_present_types]
        if missing_rec:
            missing_records[rec_id] = missing_rec

    all_present = len(missing_records) == 0

    return {
        'all_present': all_present,
        'missing_transaction_docs': [],
        'missing_record_docs': missing_records,
    }


def _resolve_names_batch(tx: Any, records: List[Any], db: Session) -> Dict[str, Any]:
    """Pre-resolve material, origin, and destination names for all records in batch."""
    from GEPPPlatform.models.users.user_location import UserLocation
    from GEPPPlatform.models.cores.references import Material

    material_ids = [r.material_id for r in records if r.material_id]
    dest_ids = [r.destination_id for r in records if r.destination_id]

    mat_map = {}
    if material_ids:
        mats = db.query(Material).filter(Material.id.in_(material_ids)).all()
        mat_map = {m.id: (m.name_en or m.name_th or '') for m in mats}

    dest_map = {}
    if dest_ids:
        dests = db.query(UserLocation).filter(UserLocation.id.in_(dest_ids)).all()
        dest_map = {d.id: (d.name_en or '') for d in dests}

    origin_name = ''
    if tx.origin_id:
        origin_loc = db.query(UserLocation).filter(UserLocation.id == tx.origin_id).first()
        origin_name = (origin_loc.name_en or '') if origin_loc else ''

    return {'mat_map': mat_map, 'dest_map': dest_map, 'origin_name': origin_name}


_BANGKOK_TZ = timezone(timedelta(hours=7))


def _build_record_data(record: Any, names: Dict[str, Any]) -> Dict[str, Any]:
    """Build a record data dict with resolved names for prompts."""
    mat_name = names['mat_map'].get(record.material_id, '') if record.material_id else ''
    dest_name = names['dest_map'].get(record.destination_id, '') if record.destination_id else ''

    # Convert transaction_date to Bangkok time (UTC+7) before formatting
    tx_date_str = None
    if record.transaction_date:
        dt = record.transaction_date
        if dt.tzinfo is not None:
            dt = dt.astimezone(_BANGKOK_TZ)
        else:
            # Assume UTC if naive
            dt = dt.replace(tzinfo=timezone.utc).astimezone(_BANGKOK_TZ)
        tx_date_str = dt.strftime('%Y-%m-%d')

    return {
        'record_id': record.id,
        'material_name': mat_name,
        'origin_name': names['origin_name'],
        'destination_name': dest_name,
        'origin_weight_kg': float(record.origin_weight_kg) if record.origin_weight_kg else 0,
        'origin_quantity': float(record.origin_quantity) if record.origin_quantity else 0,
        'origin_price_per_unit': float(record.origin_price_per_unit) if record.origin_price_per_unit else 0,
        'total_amount': float(record.total_amount) if record.total_amount else 0,
        'transaction_date': tx_date_str,
    }


def _step8a_transaction_level_check(
    image_data: Dict[str, Any],
    classified_evidence: List[Dict],
    checklist_columns: List[str],
    all_records_data: List[Dict[str, Any]],
    llm: Any,
    token_usage: Dict[str, int],
    transaction_id: int,
) -> Dict[str, Dict]:
    """
    Phase A: Check ALL records at once per tx-level evidence file.
    Returns tx_checklist: {col: {match: bool, found: bool, error: str|None}}
    """
    from ..prompts.builders.transaction_evidence_matching import build_transaction_checklist_prompt
    from ..clients.llm_client import call_llm_text_only, parse_json_response

    # Initialize checklist — all False
    tx_checklist = {
        col: {'match': False, 'found': False, 'error': None}
        for col in checklist_columns
    }

    # Get tx-level evidence files
    tx_file_ids = set(image_data.get('transaction_file_ids', []))
    tx_evidence = [
        {
            'file_id': ev['file_id'],
            'document_type_name': ev.get('document_type_name', 'unknown'),
            'extracted_data': ev.get('extracted_data', {}),
        }
        for ev in classified_evidence
        if ev['file_id'] in tx_file_ids and ev.get('extracted_data')
    ]

    if not tx_evidence or not all_records_data:
        return tx_checklist

    def _check_single_evidence(ev):
        """Check one evidence file against all records (thread-safe), with retry on JSON parse failure."""
        total_usage = {}
        prompt = build_transaction_checklist_prompt(
            single_evidence_data=ev,
            all_records_data=all_records_data,
            checklist_columns=checklist_columns,
        )
        for attempt in range(3):
            try:
                response = call_llm_text_only(llm, prompt)
                usage = response.get('usage', {})
                total_usage['input_tokens'] = total_usage.get('input_tokens', 0) + usage.get('input_tokens', 0)
                total_usage['output_tokens'] = total_usage.get('output_tokens', 0) + usage.get('output_tokens', 0)
                parsed = parse_json_response(response['content'])
                return parsed, total_usage
            except json.JSONDecodeError as je:
                if attempt < 2:
                    logger.warning(f"Tx #{transaction_id} Phase A file #{ev.get('file_id')}: JSON parse failed (attempt {attempt+1}/3), retrying...")
                    continue
                logger.error(f"Tx #{transaction_id} Phase A file #{ev.get('file_id')}: JSON parse failed after 3 attempts: {str(je)}")
                return None, total_usage
            except Exception as e:
                logger.error(f"Tx #{transaction_id} Phase A evidence file #{ev.get('file_id')} failed: {str(e)}")
                return None, total_usage
        return None, total_usage

    # Run evidence checks in parallel
    all_parsed = []
    with ThreadPoolExecutor(max_workers=len(tx_evidence)) as ev_executor:
        futures = {ev_executor.submit(_check_single_evidence, ev): ev for ev in tx_evidence}
        for future in as_completed(futures):
            parsed, usage = future.result()
            token_usage['input_tokens'] += usage.get('input_tokens', 0)
            token_usage['output_tokens'] += usage.get('output_tokens', 0)
            if parsed:
                all_parsed.append(parsed)

    # OR-merge all results
    for parsed in all_parsed:
        match_result = parsed.get('match', {})
        found_result = parsed.get('found', {})
        errors_result = parsed.get('errors', {})

        for col in checklist_columns:
            if col not in tx_checklist:
                continue
            col_match = bool(match_result.get(col, False))
            col_found = bool(found_result.get(col, False))
            col_error = errors_result.get(col)

            if col_match:
                tx_checklist[col]['match'] = True
                tx_checklist[col]['found'] = True
                tx_checklist[col]['error'] = None
            elif col_found:
                tx_checklist[col]['found'] = True
                if not tx_checklist[col]['match'] and col_error:
                    tx_checklist[col]['error'] = col_error

    return tx_checklist


def _step8b_record_level_check(
    records: List[Any],
    image_data: Dict[str, Any],
    classified_evidence: List[Dict],
    unmatched_columns: List[str],
    names: Dict[str, Any],
    llm: Any,
    token_usage: Dict[str, int],
    transaction_id: int,
) -> Dict[int, Dict[str, Dict]]:
    """
    Phase B: For each record, check record-level evidence for unmatched columns only.
    Returns per_record_results: {record_id: {col: {match, found, error}}}
    """
    from ..prompts.builders.transaction_record_evidence_matching import build_record_checklist_prompt
    from ..clients.llm_client import call_llm_text_only, parse_json_response

    per_record_results = {}
    record_file_ids = image_data.get('record_file_ids', {})
    tx_file_ids_set = set(image_data.get('transaction_file_ids', []))

    # Build classified evidence lookup by file_id
    ev_by_id = {}
    for ev in classified_evidence:
        if ev.get('extracted_data'):
            ev_by_id[ev['file_id']] = {
                'file_id': ev['file_id'],
                'document_type_name': ev.get('document_type_name', 'unknown'),
                'extracted_data': ev.get('extracted_data', {}),
            }

    # Pre-build tx-level evidence list (reusable for all records)
    tx_evidence_list = [ev_by_id[fid] for fid in tx_file_ids_set if fid in ev_by_id]

    def _check_single_record(record):
        """Check a single record's evidence (thread-safe, no DB access), with retry on JSON parse failure."""
        rec_checklist = {
            col: {'match': False, 'found': False, 'error': None}
            for col in unmatched_columns
        }

        # Each record gets its OWN files + tx-level files only.
        # A record without its own evidence should be rejected independently.
        rec_fids = record_file_ids.get(record.id, [])
        rec_evidence = [ev_by_id[fid] for fid in rec_fids if fid in ev_by_id]
        # Add tx-level evidence that isn't already in rec_evidence
        seen_fids = {fid for fid in rec_fids}
        for tx_ev in tx_evidence_list:
            if tx_ev['file_id'] not in seen_fids:
                rec_evidence.append(tx_ev)
                seen_fids.add(tx_ev['file_id'])

        if not rec_evidence:
            print(f"[AUDIT-DEBUG] Phase B record #{record.id}: NO evidence available (rec_fids={rec_fids}, tx_evidence_count={len(tx_evidence_list)})")
            return record.id, rec_checklist, {}

        total_usage = {}
        record_data = _build_record_data(record, names)
        print(f"[AUDIT-DEBUG] Phase B record #{record.id}: evidence_count={len(rec_evidence)}, evidence_file_ids={[e['file_id'] for e in rec_evidence]}, unmatched_columns={unmatched_columns}")
        print(f"[AUDIT-DEBUG] Phase B record #{record.id}: record_data={json.dumps(record_data, ensure_ascii=False)}")
        prompt = build_record_checklist_prompt(
            record_data=record_data,
            record_evidence_list=rec_evidence,
            unmatched_columns=unmatched_columns,
        )

        for attempt in range(3):
            try:
                response = call_llm_text_only(llm, prompt)
                usage = response.get('usage', {})
                total_usage['input_tokens'] = total_usage.get('input_tokens', 0) + usage.get('input_tokens', 0)
                total_usage['output_tokens'] = total_usage.get('output_tokens', 0) + usage.get('output_tokens', 0)
                parsed = parse_json_response(response['content'])

                match_result = parsed.get('match', {})
                found_result = parsed.get('found', {})
                errors_result = parsed.get('errors', {})
                print(f"[AUDIT-DEBUG] Phase B record #{record.id}: LLM response match={match_result}, found={found_result}, errors={errors_result}")

                for col in unmatched_columns:
                    rec_checklist[col] = {
                        'match': bool(match_result.get(col, False)),
                        'found': bool(found_result.get(col, False)),
                        'error': errors_result.get(col),
                    }

                return record.id, rec_checklist, total_usage

            except json.JSONDecodeError as je:
                if attempt < 2:
                    logger.warning(f"Tx #{transaction_id} Phase B record #{record.id}: JSON parse failed (attempt {attempt+1}/3), retrying...")
                    continue
                logger.error(f"Tx #{transaction_id} Phase B record #{record.id}: JSON parse failed after 3 attempts: {str(je)}")
                return record.id, rec_checklist, total_usage
            except Exception as e:
                logger.error(f"Tx #{transaction_id} Phase B record #{record.id} failed: {str(e)}")
                return record.id, rec_checklist, total_usage

        return record.id, rec_checklist, total_usage

    # Run record checks in parallel
    with ThreadPoolExecutor(max_workers=len(records)) as rec_executor:
        futures = {rec_executor.submit(_check_single_record, r): r.id for r in records}
        for future in as_completed(futures):
            try:
                rec_id, rec_checklist, usage = future.result()
                per_record_results[rec_id] = rec_checklist
                token_usage['input_tokens'] += usage.get('input_tokens', 0)
                token_usage['output_tokens'] += usage.get('output_tokens', 0)
            except Exception as e:
                rid = futures[future]
                logger.error(f"Tx #{transaction_id} Phase B record #{rid} thread failed: {str(e)}")
                per_record_results[rid] = {
                    col: {'match': False, 'found': False, 'error': None}
                    for col in unmatched_columns
                }

    return per_record_results


def _determine_final_status(
    tx_checklist: Dict[str, Dict],
    per_record_results: Dict[int, Dict[str, Dict]],
    checklist_columns: List[str],
    required_columns: List[str],
    doc_check: Dict[str, Any],
) -> Dict[str, Any]:
    """
    Phase C: Determine final status by OR-merging tx_checklist with record_aggregate.

    Returns:
        {
            'final_checklist': {col: {match, found, error}},
            'rejection_errors': [str, ...],
            'status': 'approved' | 'rejected',
        }
    """
    # Step 1: AND across all records per column → record_aggregate
    record_ids = list(per_record_results.keys())
    unmatched_columns = list(per_record_results[record_ids[0]].keys()) if record_ids else []

    record_aggregate = {}
    for col in unmatched_columns:
        # AND: column is true only if ALL records have it true
        all_match = all(
            per_record_results[rid].get(col, {}).get('match', False)
            for rid in record_ids
        ) if record_ids else False
        all_found = all(
            per_record_results[rid].get(col, {}).get('found', False)
            for rid in record_ids
        ) if record_ids else False

        # Collect first error from records where match=False
        first_error = None
        if not all_match:
            for rid in record_ids:
                err = per_record_results[rid].get(col, {}).get('error')
                if err:
                    first_error = err
                    break

        record_aggregate[col] = {
            'match': all_match,
            'found': all_found,
            'error': first_error if not all_match else None,
        }

    # Step 2: OR tx_checklist with record_aggregate → final_checklist
    final_checklist = {}
    for col in checklist_columns:
        tx = tx_checklist.get(col, {'match': False, 'found': False, 'error': None})
        rec = record_aggregate.get(col, {'match': False, 'found': False, 'error': None})

        final_match = tx['match'] or rec['match']
        final_found = tx['found'] or rec['found']
        # If final match is true, discard errors
        final_error = None if final_match else (tx['error'] or rec['error'])

        final_checklist[col] = {
            'match': final_match,
            'found': final_found,
            'error': final_error,
        }

    # Step 3: Collect rejection errors
    rejection_errors = []
    is_rejected = False

    for col in checklist_columns:
        fc = final_checklist[col]
        if fc['found'] and not fc['match']:
            # Evidence exists but data doesn't match → REJECT
            is_rejected = True
            if fc['error']:
                rejection_errors.append(fc['error'])
        elif not fc['found'] and col in required_columns:
            # Missing required evidence → REJECT
            is_rejected = True
            col_desc = _get_column_description(col)
            rejection_errors.append(f"ไม่พบข้อมูล {col_desc} ในเอกสารแนบ")

    # Also check missing docs — only reject records whose column checks also failed.
    # If a record's columns all pass (data verified from evidence), missing doc type is informational.
    missing_record_docs = doc_check.get('missing_record_docs', {})
    if missing_record_docs:
        for rec_id, missing_list in missing_record_docs.items():
            # Check if this record's columns all passed in Phase B
            rec_results = per_record_results.get(rec_id, {})
            rec_all_columns_passed = True
            if rec_results:
                for col_result in rec_results.values():
                    if not (col_result.get('match') and col_result.get('found')):
                        rec_all_columns_passed = False
                        break
            else:
                # No Phase B results — check if tx-level covers all columns
                rec_all_columns_passed = all(
                    tx_checklist.get(col, {}).get('match') and tx_checklist.get(col, {}).get('found')
                    for col in checklist_columns
                ) if checklist_columns else True

            if not rec_all_columns_passed:
                is_rejected = True
                for doc_info in missing_list:
                    doc_name = doc_info.get('name_th', doc_info.get('name_en', '')) if isinstance(doc_info, dict) else str(doc_info)
                    rejection_errors.append(f"ไม่พบเอกสาร '{doc_name}' สำหรับรายการ #{rec_id}")

    status = 'rejected' if is_rejected else 'approved'

    return {
        'final_checklist': final_checklist,
        'rejection_errors': rejection_errors,
        'status': status,
    }


def _get_column_description(col: str) -> str:
    """Get Thai description for a column name."""
    descriptions = {
        'material_id': 'ชื่อวัสดุ',
        'origin_id': 'แหล่งที่มา',
        'destination_id': 'ปลายทาง',
        'origin_weight_kg': 'น้ำหนัก (กก.)',
        'origin_quantity': 'จำนวน',
        'origin_price_per_unit': 'ราคาต่อหน่วย',
        'total_amount': 'ยอดรวม',
        'transaction_date': 'วันที่ทำรายการ',
    }
    return descriptions.get(col, col)


def _compose_audit_note(
    transaction_id: int,
    status: str,
    final_checklist: Dict[str, Dict],
    rejection_errors: List[str],
    doc_check: Dict[str, Any],
    classified_evidence: List[Dict],
) -> Dict[str, Any]:
    """Compose audit note programmatically without LLM call."""
    issues = []
    missing_record_docs = doc_check.get('missing_record_docs', {})

    # Build issues from rejection errors
    for col, result in final_checklist.items():
        if result.get('found') and not result.get('match'):
            issues.append({
                'type': 'data_mismatch',
                'field': col,
                'description': result.get('error') or f'{_get_column_description(col)} ไม่ตรงกับเอกสาร',
            })
        elif not result.get('found'):
            issues.append({
                'type': 'missing_evidence',
                'field': col,
                'description': f'ไม่พบข้อมูล {_get_column_description(col)} ในเอกสารแนบ',
            })

    for rec_id, missing_list in missing_record_docs.items():
        for doc_info in missing_list:
            doc_name = doc_info.get('name_th', doc_info.get('name_en', '')) if isinstance(doc_info, dict) else str(doc_info)
            issues.append({
                'type': 'missing_document',
                'field': doc_name,
                'description': f'ไม่พบเอกสาร \'{doc_name}\' สำหรับรายการ #{rec_id}',
            })

    n_evidence = len(classified_evidence)

    if status == 'approved':
        summary_th = f'ตรวจสอบเอกสาร {n_evidence} ไฟล์เรียบร้อย ข้อมูลตรงกับเอกสารทั้งหมด'
        summary_en = f'Verified {n_evidence} evidence files. All data matches.'
    else:
        n_issues = len(issues)
        summary_th = f'ตรวจสอบเอกสาร {n_evidence} ไฟล์ พบปัญหา {n_issues} รายการ'
        summary_en = f'Verified {n_evidence} evidence files. Found {n_issues} issue(s).'

    return {
        'status': status,
        'summary_th': summary_th,
        'summary_en': summary_en,
        'issues': issues,
    }


def _step9_compose_and_save(
    tx: Any,
    records: List[Any],
    doc_check: Dict[str, Any],
    final_determination: Dict[str, Any],
    classified_evidence: List[Dict],
    per_record_results: Dict[int, Dict[str, Dict]],
    tx_checklist: Dict[str, Dict],
    token_usage: Dict[str, int],
    processing_start: float,
    organization_id: int,
    model_version: str,
    db: Session,
    llm: Any,
    image_data: Optional[Dict[str, Any]] = None,
) -> Dict[str, Any]:
    """
    Step 9: Compose final audit note, update statuses, insert TransactionAudit record.
    Uses the new checklist-based final_determination from Phase C.
    """
    from GEPPPlatform.models.transactions.transactions import AIAuditStatus
    from GEPPPlatform.models.transactions.transaction_audits import TransactionAudit

    final_checklist = final_determination['final_checklist']
    rejection_errors = list(final_determination['rejection_errors'])
    determined_status = final_determination['status']

    # Derive checklist_columns from tx_checklist keys
    checklist_columns = list(tx_checklist.keys())

    print(f"[AUDIT-DEBUG] Tx #{tx.id} Step9 version={AUDIT_SCRIPT_VERSION}")
    print(f"[AUDIT-DEBUG] Tx #{tx.id} Step9: image_data tx_file_ids={image_data.get('transaction_file_ids', []) if image_data else 'None'}, record_file_ids={image_data.get('record_file_ids', {}) if image_data else 'None'}")
    print(f"[AUDIT-DEBUG] Tx #{tx.id} Step9: checklist_columns={checklist_columns}")
    print(f"[AUDIT-DEBUG] Tx #{tx.id} Step9: tx_checklist={tx_checklist}")
    print(f"[AUDIT-DEBUG] Tx #{tx.id} Step9: per_record_results={per_record_results}")
    print(f"[AUDIT-DEBUG] Tx #{tx.id} Step9: determined_status={determined_status}, rejection_errors={rejection_errors}")

    # Check per-record: reject records that have no evidence (no own files + no tx-level files)
    if image_data is None:
        image_data = {}
    tx_file_ids_set = set(image_data.get('transaction_file_ids', []))
    tx_has_files = any(ev['file_id'] in tx_file_ids_set for ev in classified_evidence)
    for record in records:
        record_images = record.images if hasattr(record, 'images') and record.images else []
        rec_own_files = image_data.get('record_file_ids', {}).get(record.id, [])
        if len(record_images) == 0 and len(rec_own_files) == 0 and not tx_has_files and checklist_columns:
            determined_status = 'rejected'
            rejection_errors.append(f'รายการ #{record.id} ไม่มีเอกสารแนบ')

    # Compose audit note programmatically (no LLM call needed)
    audit_note = _compose_audit_note(
        tx.id, determined_status, final_checklist, rejection_errors, doc_check, classified_evidence,
    )

    # Use the determined status from Phase C (not LLM's opinion)
    final_status = determined_status
    processing_time_ms = int((time.time() - processing_start) * 1000)

    # Update transaction
    if final_status == 'approved':
        tx.ai_audit_status = AIAuditStatus.approved
    else:
        tx.ai_audit_status = AIAuditStatus.rejected

    tx.ai_audit_note = {
        'status': final_status,
        'summary_th': audit_note.get('summary_th', ''),
        'summary_en': audit_note.get('summary_en', ''),
        'issues': audit_note.get('issues', []),
        'evidence_classified': len(classified_evidence),
        'doc_check': doc_check,
        'final_checklist': final_checklist,
        'tx_checklist': tx_checklist,
        'per_record_results': {str(k): v for k, v in per_record_results.items()},
    }
    tx.ai_audit_date = datetime.now(timezone.utc)
    tx.audit_tokens = token_usage
    flag_modified(tx, 'ai_audit_note')
    flag_modified(tx, 'audit_tokens')

    # Update each record's ai_audit_status and ai_audit_note
    record_file_ids = image_data.get('record_file_ids', {})
    for record in records:
        rec_checklist = per_record_results.get(record.id, {})
        rec_missing = doc_check.get('missing_record_docs', {}).get(record.id, [])

        # Check if the ENTIRE transaction has NO evidence files at all
        # Check if THIS record has no evidence files (neither own images nor tx-level images)
        record_images = record.images if hasattr(record, 'images') and record.images else []
        rec_own_files = record_file_ids.get(record.id, [])
        rec_has_no_files = len(record_images) == 0 and len(rec_own_files) == 0 and not tx_has_files

        print(f"[AUDIT-DEBUG] Tx #{tx.id} Record #{record.id}: rec_checklist={rec_checklist}, rec_missing={rec_missing}, rec_has_no_files={rec_has_no_files}, images={record_images}, rec_own_files={rec_own_files}, tx_has_files={tx_has_files}")

        # Per-record errors: only flag issues specific to THIS record
        # Phase B now includes tx-level evidence, so rec_checklist reflects the full picture
        rec_errors = []
        rec_has_issue = False

        # Check columns from rec_checklist (Phase B results)
        for col, result in rec_checklist.items():
            if result.get('match') and result.get('found'):
                # Evidence found and matches this record → PASS
                print(f"[AUDIT-DEBUG]   Record #{record.id} col={col}: PASS (found=T, match=T)")
            elif result.get('found') and not result.get('match'):
                # Evidence found but doesn't match this record → REJECT
                rec_has_issue = True
                col_desc = _get_column_description(col)
                err_msg = result.get('error') or f'{col_desc} ไม่ตรงกับเอกสาร'
                rec_errors.append(err_msg)
                print(f"[AUDIT-DEBUG]   Record #{record.id} col={col}: REJECT (found=T, match=F, error={err_msg})")
            elif not result.get('found'):
                # No evidence found for this column at all (neither tx nor record level)
                tx_col = tx_checklist.get(col, {})
                if tx_col.get('match') and tx_col.get('found'):
                    # Tx-level Phase A already matched all records → PASS
                    print(f"[AUDIT-DEBUG]   Record #{record.id} col={col}: PASS (tx-level matched all)")
                else:
                    # No evidence anywhere → missing evidence
                    rec_has_issue = True
                    col_desc = _get_column_description(col)
                    rec_errors.append(f'ไม่พบข้อมูล {col_desc} ในเอกสารแนบ')
                    print(f"[AUDIT-DEBUG]   Record #{record.id} col={col}: REJECT (no evidence anywhere)")

        # Also check columns that were NOT in rec_checklist (Phase B skipped or returned empty)
        # but ARE required — these need tx-level Phase A pass or they fail
        for col in checklist_columns:
            if col in rec_checklist:
                continue  # Already handled above
            tx_col = tx_checklist.get(col, {})
            if tx_col.get('match') and tx_col.get('found'):
                # Tx-level Phase A matched all records → PASS
                print(f"[AUDIT-DEBUG]   Record #{record.id} col={col}: PASS (tx-level matched, not in rec_checklist)")
            else:
                # Required column not verified by any phase
                rec_has_issue = True
                col_desc = _get_column_description(col)
                rec_errors.append(f'ไม่พบข้อมูล {col_desc} ในเอกสารแนบ')
                print(f"[AUDIT-DEBUG]   Record #{record.id} col={col}: REJECT (not in rec_checklist, tx-level not matched)")

        # Record with no evidence files must be rejected (when there are required columns)
        if rec_has_no_files and checklist_columns:
            rec_has_issue = True
            rec_errors.append('ไม่มีเอกสารแนบสำหรับรายการนี้')
            print(f"[AUDIT-DEBUG]   Record #{record.id}: REJECT (no files at all)")

        # Safety net: if rec_has_issue but no errors were added, add a generic message
        if rec_has_issue and not rec_errors:
            rec_errors.append('ข้อมูลในเอกสารแนบไม่ตรงกับรายการ')
            print(f"[AUDIT-DEBUG]   Record #{record.id}: added fallback error (rec_has_issue=True but rec_errors was empty)")

        # Determine per-record status independently:
        # - Column matching (rec_has_issue) is the primary check
        # - Missing doc types (rec_missing) only reject if column checks ALSO failed
        #   or if the record has no evidence at all. If all required columns matched
        #   from the evidence, the record's data is verified regardless of doc type classification.
        rec_columns_all_passed = not rec_has_issue and checklist_columns
        if rec_missing and rec_columns_all_passed:
            # Record has missing doc types BUT all column checks passed from evidence
            # → data is verified, doc type mismatch is informational only
            print(f"[AUDIT-DEBUG]   Record #{record.id}: missing docs={rec_missing} but all columns passed → APPROVED (data verified)")
            rec_missing_for_status = []
        else:
            rec_missing_for_status = rec_missing

        print(f"[AUDIT-DEBUG]   Record #{record.id}: final rec_has_issue={rec_has_issue}, rec_missing={rec_missing}, rec_missing_for_status={rec_missing_for_status}")

        if rec_has_issue or rec_missing_for_status:
            record.ai_audit_status = 'rejected'
        else:
            # Record has no individual issues — approve it even if tx-level is rejected
            # (other records may have caused the tx-level rejection)
            record.ai_audit_status = 'approved'

        record.ai_audit_note = {
            'status': record.ai_audit_status,
            'checklist': rec_checklist,
            'errors': rec_errors,
            'missing_docs': rec_missing,  # Keep original for informational display
        }
        flag_modified(record, 'ai_audit_note')
        print(f"[AUDIT-DEBUG]   Record #{record.id}: STATUS={record.ai_audit_status}")

    # Insert TransactionAudit record with status snapshots
    tx_status_value = tx.status.value if hasattr(tx.status, 'value') else str(tx.status)
    audit_record = TransactionAudit(
        transaction_id=tx.id,
        audit_notes=tx.ai_audit_note,
        by_human=False,
        auditor_id=None,
        organization_id=organization_id,
        audit_type='ai_not_sync',
        audit_status=tx_status_value,
        ai_audit_status=final_status,
        processing_time_ms=processing_time_ms,
        token_usage=token_usage,
        model_version=model_version,
    )
    db.add(audit_record)
    db.flush()  # Get the audit_record.id before commit

    logger.info(f"Transaction #{tx.id}: audit {final_status} (processing: {processing_time_ms}ms, tokens: {token_usage})")

    return {
        'transaction_id': tx.id,
        'status': final_status,
        'processing_time_ms': processing_time_ms,
        'token_usage': token_usage,
        'audit_record_id': audit_record.id,
    }
